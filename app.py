"""
AI-Powered Resume & Cover Letter Generator
-----------------------------------------
A Streamlit app where users upload a resume (PDF/DOCX) or paste their LinkedIn profile,
paste a job description, and receive a tailored resume version and a customized cover letter.

Features
- Upload resume (PDF/DOCX) or paste LinkedIn/profile text
- Paste job description text or upload JD file (PDF/TXT)
- Tone, industry, seniority, and role selectors
- Team presets (upload JSON) to fine-tune behavior by industry/role
- One-click: Analyze Fit, Tailor Resume, Generate Cover Letter
- Download as DOCX or Markdown
- Optional anonymization / PII scrubbing

How to run locally
------------------
1) Save this file as `app.py`.
2) Create and activate a virtual environment (optional but recommended).
3) `pip install -r requirements.txt` (see requirements block at bottom of this file)
4) Set your OpenAI API key in the environment: `export OPENAI_API_KEY=your_key`
5) `streamlit run app.py`

Note: For Google AI Studio (Gemini) support, set `PROVIDER = "google"` and
provide `GOOGLE_API_KEY`. By default this file uses OpenAI. Both providers
share the same UI.
"""

import io
import json
import os
import openai
import re
import time
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple

import streamlit as st

# -------- Optional LLM providers --------
openai.api_key = os.getenv("OPENAI_API_KEY")
PROVIDER = os.getenv("LLM_PROVIDER", "openai").lower()  # "openai" or "google"
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4.1-mini")
GOOGLE_MODEL = os.getenv("GOOGLE_MODEL", "gemini-2.0-flash-exp")

# Lazy imports so the app can start even if a provider isn't installed
_openai_client = None
_google_client = None


def get_openai_client():
    global _openai_client
    if _openai_client is None:
        try:
            from openai import OpenAI
            _openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        except Exception as e:
            st.error(f"OpenAI client not available: {e}")
    return _openai_client


def get_google_client():
    global _google_client
    if _google_client is None:
        try:
            import google.generativeai as genai
            genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
            _google_client = genai.GenerativeModel(GOOGLE_MODEL)
        except Exception as e:
            st.error(f"Google Generative AI not available: {e}")
    return _google_client


# --------- File helpers ---------

def extract_text_from_pdf(uploaded_file) -> str:
    """Extracts text from a PDF using pdfminer.six."""
    try:
        from pdfminer.high_level import extract_text
        with io.BytesIO(uploaded_file.read()) as f:
            return extract_text(f) or ""
    except Exception as e:
        st.warning(f"PDF text extraction failed: {e}")
        return ""


def extract_text_from_docx(uploaded_file) -> str:
    try:
        import docx
        doc = docx.Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        st.warning(f"DOCX text extraction failed: {e}")
        return ""


def read_text_file(uploaded_file) -> str:
    try:
        return uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception:
        return ""


# --------- Domain presets / Team profiles ---------
DEFAULT_PRESETS = {
    "Software Engineering": {
        "tone": "confident",
        "style": "ATS-friendly, bullet-first, quantitative impact",
        "keywords": [
            "Python", "Java", "C++", "System Design", "Microservices", "REST", "GraphQL",
            "Docker", "Kubernetes", "AWS", "GCP", "CI/CD", "Unit Testing", "Agile", "Scrum",
        ],
    },
    "Data Science": {
        "tone": "analytical",
        "style": "Results-focused, highlights experiments, metrics and business impact",
        "keywords": [
            "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "PyTorch", "A/B Testing",
            "Causal Inference", "SQL", "Airflow", "ML Ops", "LLMs", "RAG",
        ],
    },
    "Product Management": {
        "tone": "strategic",
        "style": "Customer-centric, combines discovery and delivery outcomes",
        "keywords": [
            "Roadmapping", "OKRs", "Discovery", "User Research", "A/B Testing", "Analytics",
            "Stakeholder Management", "GTM", "Monetization",
        ],
    },
    "Sales": {
        "tone": "persuasive",
        "style": "Quota-crushing, pipeline management, enterprise wins",
        "keywords": [
            "Prospecting", "CRM", "Salesforce", "Enterprise Sales", "Negotiation", "Forecasting",
            "Solution Selling", "MEDDIC",
        ],
    },
}


# --------- Prompting ---------
BASE_SYSTEM_PROMPT = (
    "You are a senior career coach and ATS optimization expert. "
    "Given a candidate profile and a target job description, you will: "
    "(1) Summarize the candidate's core strengths in the context of the role, "
    "(2) Draft a tailored, ATS-friendly resume section-by-section, emphasizing quantifiable impact, "
    "(3) Draft a concise cover letter with the requested tone, "
    "(4) Ensure the content reflects regional spelling for the selected locale, avoids exaggeration, and remains truthful, "
    "(5) Mirror keywords from the job description where appropriate without keyword stuffing, and "
    "(6) Keep formatting simple (headings, bullet points, no tables)."
)

RESUME_TEMPLATE = (
    "# Tailored Resume\n"
    "## Header\n"
    "<NAME> | <CITY, COUNTRY> | <EMAIL> | <PHONE> | <LINKEDIN/PORTFOLIO>\n\n"
    "## Professional Summary\n"
    "${summary}\n\n"
    "## Core Skills\n"
    "- ${skills}\n\n"
    "## Experience\n"
    "${experience}\n\n"
    "## Education\n"
    "${education}\n\n"
    "## Certifications (optional)\n"
    "${certifications}\n\n"
    "## Projects (optional)\n"
    "${projects}\n"
)

COVER_LETTER_TEMPLATE = (
    "# Tailored Cover Letter\n\n"
    "${greeting}\n\n"
    "${opening}\n\n"
    "${body}\n\n"
    "${closing}\n\n"
    "Sincerely,\n\n<NAME>"
)


@dataclass
class GenerationRequest:
    candidate_text: str
    job_text: str
    tone: str
    industry: str
    role: str
    seniority: str
    locale: str
    extra_keywords: List[str]
    style_overrides: Optional[str] = None
    anonymize: bool = False


def build_user_prompt(req: GenerationRequest) -> str:
    keyline = ", ".join(req.extra_keywords) if req.extra_keywords else ""
    return f"""
You will create two outputs for the candidate:
1) A tailored resume that follows the RESUME TEMPLATE sections and is ATS-friendly.
2) A tailored cover letter that follows the COVER LETTER TEMPLATE and the tone requested.

Context & Constraints:
- Industry: {req.industry}
- Target Role: {req.role}
- Seniority: {req.seniority}
- Locale/Spelling: {req.locale}
- Desired Tone: {req.tone}
- Extra keywords to weave in: {keyline}
- If information is missing, infer sensibly but mark with <PLACEHOLDER>.
- Avoid tables or images.
- Use short, high-impact bullets with quantified outcomes (e.g., 23% growth, $1.2M ARR).
- Mirror language from the job description without copying verbatim.
- Respect truthfulness: never invent employers or degrees; suggest placeholders instead.
- Provide both artifacts in GitHub-flavored Markdown.
- If anonymize=True, remove PII like full name, phone, personal email; use <REDACTED>.

Candidate Profile (Resume/LinkedIn):
"""
    + req.candidate_text.strip() + "\n\n" + "Job Description:\n" + req.job_text.strip() + "\n\n" + "Output:"


# --------- LLM call wrappers ---------

def call_openai(system: str, user: str) -> str:
    client = get_openai_client()
    if not client:
        raise RuntimeError("OpenAI client not initialized. Set OPENAI_API_KEY.")
    try:
        # Prefer the Responses API when available
        resp = client.responses.create(
            model=OPENAI_MODEL,
            input=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
        )
        return resp.output_text
    except Exception:
        # Fallback: Chat Completions (for older models)
        try:
            chat = client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
                temperature=0.3,
            )
            return chat.choices[0].message.content
        except Exception as e:
            raise RuntimeError(f"OpenAI generation failed: {e}")


def call_google(system: str, user: str) -> str:
    model = get_google_client()
    if not model:
        raise RuntimeError("Google Generative AI client not initialized. Set GOOGLE_API_KEY.")
    try:
        full_prompt = f"System Instructions:\n{system}\n\nUser:\n{user}"
        resp = model.generate_content(full_prompt)
        return resp.text or ""
    except Exception as e:
        raise RuntimeError(f"Google generation failed: {e}")


def generate_outputs(req: GenerationRequest) -> str:
    system = BASE_SYSTEM_PROMPT
    user = build_user_prompt(req)
    if PROVIDER == "google":
        return call_google(system, user)
    return call_openai(system, user)


# --------- Rendering / Export ---------

def split_outputs(md_text: str) -> Tuple[str, str]:
    """Heuristically split markdown into resume and cover letter blocks."""
    # Look for headings
    resume_idx = re.search(r"#\s*Tailored Resume", md_text, flags=re.I)
    cover_idx = re.search(r"#\s*Tailored Cover Letter", md_text, flags=re.I)
    if resume_idx and cover_idx:
        if resume_idx.start() < cover_idx.start():
            return md_text[resume_idx.start():cover_idx.start()].strip(), md_text[cover_idx.start():].strip()
        else:
            return md_text[cover_idx.start():resume_idx.start()].strip(), md_text[resume_idx.start():].strip()
    # Fallback: naive split by first big heading
    parts = re.split(r"\n#\s+", md_text, maxsplit=1)
    if len(parts) == 2:
        first = parts[0].strip()
        second = parts[1].strip()
        if "cover" in second.lower():
            return first, "# " + second
        return "# " + second, first
    return md_text.strip(), ""


def to_docx(md_text: str) -> bytes:
    """Very simple Markdown-to-DOCX by stripping markdown and writing paragraphs."""
    try:
        import docx
    except Exception:
        st.warning("python-docx not installed; cannot export DOCX.")
        return b""

    # Basic cleanup: remove markdown symbols but keep bullets
    plain = re.sub(r"^#.*$", "", md_text, flags=re.M)
    plain = re.sub(r"\*\*(.*?)\*\*", r"\1", plain)
    plain = re.sub(r"_(.*?)_", r"\1", plain)

    doc = docx.Document()
    for line in plain.splitlines():
        if line.strip().startswith(('-', '*')):
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(line.strip().lstrip('-* ').strip())
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# --------- UI ---------
st.set_page_config(page_title="AI Resume & Cover Letter Generator", page_icon="🧰", layout="wide")

st.title("🧰 AI-Powered Resume & Cover Letter Generator")
st.caption("Upload your resume or paste LinkedIn text, add a job description, and get tailored outputs.")

with st.sidebar:
    st.header("Settings")
    st.write("**Provider**: set via `LLM_PROVIDER` env var (openai/google)")

    if PROVIDER == "google":
        st.text_input("GOOGLE_API_KEY", type="password", value=os.getenv("GOOGLE_API_KEY", ""))
    else:
        st.text_input("OPENAI_API_KEY", type="password", value=os.getenv("OPENAI_API_KEY", ""))

    model_name = st.text_input(
        "Model name",
        value=GOOGLE_MODEL if PROVIDER == "google" else OPENAI_MODEL,
        help="Override the default model if desired.",
    )
    if PROVIDER == "google":
        os.environ["GOOGLE_MODEL"] = model_name
    else:
        os.environ["OPENAI_MODEL"] = model_name

    st.divider()
    st.subheader("Team Presets")
    preset_choice = st.selectbox("Select industry preset", list(DEFAULT_PRESETS.keys()))
    uploaded_preset = st.file_uploader("Or upload a JSON team preset", type=["json"], accept_multiple_files=False)

    presets = DEFAULT_PRESETS.copy()
    if uploaded_preset is not None:
        try:
            custom = json.loads(uploaded_preset.read().decode("utf-8"))
            # merge (uploaded overrides defaults on matching keys)
            presets.update(custom)
        except Exception as e:
            st.warning(f"Failed to parse preset JSON: {e}")

    selected_preset = presets.get(preset_choice, {})

st.subheader("1) Candidate Profile")
col1, col2 = st.columns(2)
with col1:
    resume_file = st.file_uploader("Upload Resume (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])
with col2:
    linked_text = st.text_area(
        "Or paste LinkedIn / profile text",
        height=180,
        placeholder="Paste your LinkedIn 'About', Experience, Skills sections, or any profile text...",
    )

candidate_text = ""
if resume_file is not None:
    if resume_file.name.lower().endswith('.pdf'):
        candidate_text = extract_text_from_pdf(resume_file)
    elif resume_file.name.lower().endswith('.docx'):
        candidate_text = extract_text_from_docx(resume_file)
    else:
        candidate_text = read_text_file(resume_file)

if not candidate_text and linked_text:
    candidate_text = linked_text

st.subheader("2) Job Description")
jd_col1, jd_col2 = st.columns(2)
with jd_col1:
    job_desc = st.text_area("Paste job description", height=220)
with jd_col2:
    jd_file = st.file_uploader("Or upload JD (PDF/TXT)", type=["pdf", "txt"])
    if jd_file is not None and not job_desc:
        if jd_file.name.lower().endswith('.pdf'):
            job_desc = extract_text_from_pdf(jd_file)
        else:
            job_desc = read_text_file(jd_file)

st.subheader("3) Controls")
cc1, cc2, cc3, cc4 = st.columns(4)
with cc1:
    tone = st.selectbox("Tone", [
        "professional", "confident", "friendly", "enthusiastic", "formal", "concise",
    ], index=0 if not selected_preset else 0)
with cc2:
    industry = st.text_input("Industry", value=preset_choice)
with cc3:
    role = st.text_input("Target Role", value="Software Engineer")
with cc4:
    seniority = st.selectbox("Seniority", ["Intern", "Junior", "Mid", "Senior", "Lead", "Manager"], index=2)

lc1, lc2, lc3 = st.columns(3)
with lc1:
    locale = st.selectbox("Locale", ["US", "UK", "India", "EU"], index=2)
with lc2:
    extra_keywords = st.text_input("Extra keywords (comma-separated)", value=", ".join(selected_preset.get("keywords", [])))
with lc3:
    anonymize = st.checkbox("Anonymize PII", value=False)

style_overrides = st.text_area(
    "Style overrides (optional)",
    value=selected_preset.get("style", ""),
    placeholder="e.g., Use strong action verbs, 1-2 lines per bullet, include metrics, ATS-friendly formatting.",
)

st.divider()

# Action buttons
b1, b2, b3 = st.columns([1,1,1])
with b1:
    analyze = st.button("Analyze Fit")
with b2:
    tailor_resume = st.button("Tailor Resume")
with b3:
    gen_cover = st.button("Generate Cover Letter")

if not candidate_text:
    st.info("Upload a resume or paste profile text to begin.")
if not job_desc:
    st.info("Paste or upload a job description to tailor outputs.")

# Perform generation if requested
generated_md = ""
if (analyze or tailor_resume or gen_cover) and candidate_text and job_desc:
    with st.spinner("Generating with LLM..."):
        req = GenerationRequest(
            candidate_text=candidate_text,
            job_text=job_desc,
            tone=tone,
            industry=industry,
            role=role,
            seniority=seniority,
            locale=locale,
            extra_keywords=[k.strip() for k in extra_keywords.split(',') if k.strip()],
            style_overrides=style_overrides,
            anonymize=anonymize,
        )
        try:
            generated_md = generate_outputs(req)
        except Exception as e:
            st.error(str(e))

if generated_md:
    resume_md, cover_md = split_outputs(generated_md)

    st.subheader("🔎 Fit Analysis & Drafts")
    tabs = st.tabs(["Tailored Resume", "Cover Letter", "Raw Markdown"])
    with tabs[0]:
        st.markdown(resume_md)
        r_docx = to_docx(resume_md)
        st.download_button(
            label="Download Resume (DOCX)",
            data=r_docx,
            file_name="tailored_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            disabled=not bool(r_docx),
        )
        st.download_button(
            label="Download Resume (Markdown)",
            data=resume_md.encode("utf-8"),
            file_name="tailored_resume.md",
            mime="text/markdown",
        )
    with tabs[1]:
        st.markdown(cover_md)
        c_docx = to_docx(cover_md)
        st.download_button(
            label="Download Cover Letter (DOCX)",
            data=c_docx,
            file_name="cover_letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            disabled=not bool(c_docx),
        )
        st.download_button(
            label="Download Cover Letter (Markdown)",
            data=cover_md.encode("utf-8"),
            file_name="cover_letter.md",
            mime="text/markdown",
        )
    with tabs[2]:
        st.code(generated_md, language="markdown")

st.divider()
with st.expander("📦 requirements.txt (copy into a file and `pip install -r requirements.txt`)"):
    st.code(
        """
streamlit>=1.36.0
openai>=1.40.0
pdfminer.six>=20231228
python-docx>=1.1.0
# Optional for Google AI Studio
google-generativeai>=0.8.3
        """.strip(),
        language="text",
    )

with st.expander("🧪 Example Team Presets JSON (upload in sidebar to override defaults)"):
    st.code(
        json.dumps(
            {
                "FinTech Engineering": {
                    "tone": "professional",
                    "style": "Security-first, regulated env, latency & reliability metrics",
                    "keywords": ["KYC", "PCI-DSS", "Latency", "Throughput", "Kafka", "Event-driven"],
                },
                "Healthcare Data Science": {
                    "tone": "empathetic",
                    "style": "HIPAA, PHI handling, outcome metrics, clinical collaboration",
                    "keywords": ["FHIR", "HL7", "PHI", "De-identification", "Survival Analysis"],
                },
            },
            indent=2,
        ),
        language="json",
    )

with st.expander("🔐 Notes on Privacy & PII"):
    st.markdown(
        "- All processing happens through the selected LLM provider; do not paste sensitive data you cannot share.\n"
        "- Use the **Anonymize PII** toggle to scrub names, emails, and phone numbers in outputs.\n"
        "- Review drafts for accuracy before sending to employers.\n"
    )


